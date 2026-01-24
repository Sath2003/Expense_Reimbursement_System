from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
from sqlalchemy.orm import Session
from app.database import get_db
from app.schemas.expense import (
    ExpenseCreate, ExpenseUpdate, ExpenseResponse, 
    ExpenseDetailResponse, ExpenseWithAttachments
)
from app.services.expense_service import ExpenseService
from app.services.approval_service import ApprovalService
from app.services.notification_service import NotificationService
from app.utils.file_handler import FileHandler
from app.utils.receipt_extractor import ReceiptExtractor
from app.utils.improved_receipt_extractor import ImprovedReceiptExtractor
from app.services.receipt_validation_service import ReceiptValidationService
from app.services.expense_cross_check_service import ExpenseCrossCheckService
from app.services.llm_receipt_agent import LLMReceiptAgent
from app.services.policy_service import PolicyService
from app.utils.audit_logger import AuditLogger
from app.utils.dependencies import get_current_user
from app.models.expense import Expense, ExpenseAttachment
from app.models.user import User
from typing import List
from datetime import datetime
import os
from decimal import Decimal
import os
import tempfile
import logging
import hashlib

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/expenses", tags=["expenses"])

@router.get("/policy/check")
async def check_policy(
    category_id: int,
    amount: float,
    date: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Check if an expense complies with policy before submission
    """
    try:
        from datetime import datetime
        expense_date_obj = datetime.strptime(date, '%Y-%m-%d').date()
    except ValueError:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid date format. Use YYYY-MM-DD"
        )
    
    policy_result = PolicyService.check_expense_policy(
        db=db,
        user=current_user,
        category_id=category_id,
        amount=Decimal(str(amount)),
        expense_date=expense_date_obj,
        transport_type_id=None
    )
    
    return {
        "is_compliant": policy_result.is_compliant,
        "violations": policy_result.violations,
        "allowed_amount": float(policy_result.allowed_amount) if policy_result.allowed_amount else None,
        "policy_details": policy_result.policy_details
    }

@router.get("/policy/user")
async def get_user_policies(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Get all applicable policies for the current user
    """
    policies = PolicyService.get_user_policies(db, current_user)
    return {"policies": policies}

@router.post("/submit")
async def submit_expense(
    category: str = Form(...),
    description: str = Form(...),
    date: str = Form(...),
    amount: str = Form(None),
    receipt: UploadFile = File(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Submit a new expense with optional file upload
    
    Args:
        category: Expense category (Travel, Equipment, Accommodation, Office Supplies, Meals, Other)
        description: Detailed description of the expense
        date: Date when the expense was incurred (YYYY-MM-DD format)
        amount: Amount in INR (optional if receipt is provided - will be extracted from receipt)
        receipt: Receipt file (optional if amount is provided)
    """
    try:
        print(f"Received: category={category}, description={description}, date={date}, amount={amount}, receipt={receipt}")
        
        # Validate required fields
        if not category or not description or not date:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Missing required fields. Got: category={category}, description={description}, date={date}"
            )
        
        # Amount validation: required if no receipt, optional if receipt provided
        if not amount and not receipt:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Amount is required when no receipt is provided"
            )
        
        # Map category name to category_id (aligns with frontend categoryMap)
        # Strip whitespace from category name for robustness
        category = category.strip() if category else ""
        
        category_map = {
            'Travel': 1,
            'Food': 2,
            'Food & Meals': 2,  # Alternative name
            'Accommodation': 3,
            'Office Supplies': 4,
            'Communication': 5,
            'Miscellaneous': 6,
            'Equipment': 7,
            'Meals': 8,
            'Meals & Drinks': 8,  # Alternative name
            'Other': 9,
            'Fuel': 10,
        }
        
        category_id = category_map.get(category)
        if not category_id:
            # Log available categories for debugging
            print(f"[DEBUG] Invalid category received: '{category}'")
            print(f"[DEBUG] Available categories: {list(category_map.keys())}")
            print(f"[DEBUG] Category length: {len(category)}, repr: {repr(category)}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Invalid category: {category}. Valid options are: {', '.join(category_map.keys())}"
            )
        
        # Handle amount:
        # - If amount is provided, use it
        # - If receipt is provided but no amount, try to extract from receipt
        # - If extraction fails, use 0 as placeholder and let user update manually
        # - If neither amount nor receipt, error (already validated above)
        extracted_amount = None
        extraction_note = ""
        extraction_confidence = "none"
        
        parsed_amount = None
        if amount is not None and str(amount).strip() != "":
            try:
                parsed_amount = float(amount)
            except (ValueError, TypeError):
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Invalid amount provided"
                )

        if parsed_amount is not None and parsed_amount > 0:
            amount_value = parsed_amount
        else:
            # Receipt is provided (or amount is 0/blank) - try to extract amount from it
            if receipt and receipt.filename:
                try:
                    # Save receipt temporarily to extract amount
                    file_extension = receipt.filename.split('.')[-1].lower()
                    
                    # Create temp file
                    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as tmp_file:
                        content = await receipt.read()
                        tmp_file.write(content)
                        tmp_path = tmp_file.name
                        await receipt.seek(0)  # Reset file pointer for later use
                    
                    # Try to extract amount using improved extractor
                    extracted_amount, confidence, extraction_note = ImprovedReceiptExtractor.extract_amount(tmp_path, file_extension)
                    extraction_confidence = confidence
                    
                    # Perform receipt validation
                    validation_service = ReceiptValidationService()
                    
                    # Get full text for validation
                    if file_extension == 'pdf':
                        try:
                            import pdfplumber
                            with pdfplumber.open(tmp_path) as pdf:
                                full_text = ""
                                for page in pdf.pages:
                                    full_text += (page.extract_text() or "") + "\n"
                        except:
                            full_text = extraction_note
                    elif file_extension in ['docx', 'doc']:
                        try:
                            from docx import Document
                            doc = Document(tmp_path)
                            full_text = ""
                            for paragraph in doc.paragraphs:
                                full_text += paragraph.text + " "
                            for table in doc.tables:
                                for row in table.rows:
                                    for cell in row.cells:
                                        full_text += cell.text + " "
                        except:
                            full_text = extraction_note
                    else:
                        full_text = extraction_note

                    # LLM agent validation (free local Ollama if enabled)
                    llm_result = await LLMReceiptAgent.evaluate_receipt(
                        extracted_text=full_text,
                        amount=float(extracted_amount or 0),
                        category=category,
                        description=description,
                        expense_date=date,
                    )
                    print(f"[LLM] decision={llm_result.get('decision')} risk={llm_result.get('risk_level')} reasons={llm_result.get('reasons')}")

                    # Hard gate: block on explicit block, and block on review if strict mode
                    from app.config import settings
                    if llm_result.get('decision') == 'block' or (settings.OLLAMA_STRICT and llm_result.get('decision') == 'review'):
                        raise HTTPException(
                            status_code=status.HTTP_400_BAD_REQUEST,
                            detail={
                                "error": "Receipt rejected by AI",
                                "ai": llm_result,
                            }
                        )
                    
                    # Validate receipt genuineness
                    validation_results = validation_service.validate_receipt(
                        tmp_path, full_text, extracted_amount or 0
                    )
                    
                    print(f"[VALIDATION] {validation_service.get_validation_summary(validation_results)}")
                    
                    # Perform AI cross-checking
                    cross_check_service = ExpenseCrossCheckService()
                    cross_check_results = cross_check_service.cross_check_expense(
                        file_path=tmp_path,
                        extracted_text=full_text,
                        amount=extracted_amount or 0,
                        category=category,
                        description=description,
                        date=date,
                        user_id=current_user.id,
                        db=db
                    )
                    
                    print(f"[CROSS-CHECK] {cross_check_service.get_validation_summary(cross_check_results)}")
                    
                    # Combine validation results
                    overall_confidence = (validation_results['confidence_score'] + cross_check_results['confidence_score']) / 2
                    is_approved = validation_results['is_genuine'] and cross_check_results['is_approved']
                    
                    # Combine risk factors and recommendations
                    all_risk_factors = validation_results['risk_factors'] + cross_check_results['risk_factors']
                    all_recommendations = validation_results['recommendations'] + cross_check_results['recommendations']
                    
                    # If cross-check fails, reject the submission
                    if not cross_check_results['is_approved']:
                        raise HTTPException(
                            status_code=status.HTTP_400_BAD_REQUEST,
                            detail={
                                "error": "Expense validation failed",
                                "message": f"AI cross-check detected issues: {cross_check_service.get_validation_summary(cross_check_results)}",
                                "risk_factors": all_risk_factors,
                                "recommendations": all_recommendations,
                                "confidence_score": overall_confidence
                            }
                        )
                    
                    # Clean up temp file
                    if os.path.exists(tmp_path):
                        os.remove(tmp_path)
                    
                    # Prepare AI validation data
                    ai_validation_data = {
                        'file_hash': cross_check_results.get('file_hash', ''),
                        'extracted_text_hash': cross_check_results.get('text_hash', ''),
                        'validation_score': overall_confidence,
                        'is_ai_validated': cross_check_results['is_approved'],
                        'risk_factors': cross_check_results['risk_factors'][:500],  # Limit length
                        'validation_timestamp': datetime.now().isoformat()
                    }
                    
                    if extracted_amount and extracted_amount > 0:
                        amount_value = extracted_amount
                        print(f"[SUCCESS] Extracted amount ₹{extracted_amount:.2f} from receipt ({confidence} confidence): {extraction_note}")
                    else:
                        # Extraction failed or returned 0 - use placeholder 0, user will update manually
                        amount_value = 0
                        print(f"[INFO] Could not extract amount from receipt: {extraction_note}. Using 0 as placeholder.")

                except HTTPException:
                    # Do not swallow validation failures (duplicate/fraud/etc)
                    raise
                except Exception as extract_err:
                    # If extraction fails, use placeholder amount 0
                    print(f"[ERROR] Receipt extraction error: {extract_err}")
                    amount_value = 0
            else:
                # No receipt and no amount - should not reach here (validation above)
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Amount is required when no receipt is provided"
                )
        
        expense_data = ExpenseCreate(
            amount=amount_value,
            category_id=category_id,
            description=description,
            expense_date=date  # This is the date the expense was incurred
        )
        
        # Policy Enforcement Check
        try:
            expense_date_obj = datetime.strptime(date, '%Y-%m-%d').date()
        except ValueError:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Invalid date format. Use YYYY-MM-DD"
            )

        today = datetime.utcnow().date()
        if expense_date_obj > today:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Expense date cannot be in the future"
            )

        if (today - expense_date_obj).days > 31:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Expense date is older than 1 month. Please submit within 31 days of the bill date."
            )
        
        policy_result = PolicyService.check_expense_policy(
            db=db,
            user=current_user,
            category_id=category_id,
            amount=Decimal(str(amount_value)),
            expense_date=expense_date_obj,
            transport_type_id=None  # TODO: Add transport_type_id if needed
        )
        
        if not policy_result.is_compliant:
            # Store policy violations in expense record
            policy_check_result = {
                "is_compliant": False,
                "violations": policy_result.violations,
                "allowed_amount": float(policy_result.allowed_amount) if policy_result.allowed_amount else None,
                "policy_details": policy_result.policy_details
            }
            # We'll still allow submission but mark as POLICY_EXCEPTION
            expense_data.policy_check_result = policy_check_result
            print(f"[POLICY] Violations: {policy_result.violations}")
        else:
            policy_check_result = {
                "is_compliant": True,
                "violations": [],
                "policy_details": policy_result.policy_details
            }
            expense_data.policy_check_result = policy_check_result
        
        # Create expense
        expense, error = ExpenseService.create_expense(
            db=db,
            user_id=current_user.id,
            expense_data=expense_data,
            performed_by=current_user.id,
            ai_validation_data=ai_validation_data if 'ai_validation_data' in locals() else None
        )
        
        if error:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=error
            )
        
        # Handle file upload if provided
        if receipt and receipt.filename:
            try:
                await receipt.seek(0)
                file_content = await receipt.read()
                file_hash = hashlib.sha256(file_content).hexdigest()
                # Check for duplicate receipt by hash
                existing_attachment = db.query(ExpenseAttachment).filter(ExpenseAttachment.file_hash == file_hash).first()
                if existing_attachment:
                    raise HTTPException(
                        status_code=status.HTTP_409_CONFLICT,
                        detail="Duplicate receipt detected. This file has already been submitted."
                    )
                # Reset file pointer for saving
                await receipt.seek(0)
                file_path, saved_hash, file_size, file_type = await FileHandler.save_file(
                    file=receipt,
                    expense_id=expense.id
                )
                # Ensure saved hash matches computed hash; if not, use computed
                final_hash = saved_hash if saved_hash else file_hash
                await receipt.seek(0)
                
                # Link attachment to expense
                expense_attachment = ExpenseAttachment(
                    expense_id=expense.id,
                    file_path=file_path,
                    file_name=receipt.filename,
                    file_type=file_type,
                    file_size=file_size,
                    file_hash=final_hash
                )
                db.add(expense_attachment)
                db.commit()
                print(f"Receipt saved for expense {expense.id}: {receipt.filename}")

                # If amount is still placeholder (0), attempt extraction from the saved file
                try:
                    current_amount = float(expense.amount or 0)
                except Exception:
                    current_amount = 0

                if current_amount == 0:
                    extracted_after_save, conf_after_save, note_after_save = ImprovedReceiptExtractor.extract_amount(
                        file_path,
                        file_type
                    )
                    if extracted_after_save and extracted_after_save > 0:
                        old_amount = expense.amount
                        expense.amount = extracted_after_save
                        expense.updated_at = datetime.utcnow()
                        db.commit()
                        db.refresh(expense)

                        AuditLogger.log(
                            db=db,
                            entity_type="expense",
                            entity_id=expense.id,
                            action="amount_extracted_from_receipt",
                            performed_by=current_user.id,
                            old_value={"amount": str(old_amount)},
                            new_value={"amount": str(extracted_after_save), "confidence": conf_after_save, "note": note_after_save}
                        )
            except HTTPException:
                raise
            except Exception as file_err:
                # Log error but don't fail the expense submission
                print(f"File upload error: {file_err}")
                pass
        
        # Create approval records - for both manager and HR
        ApprovalService.submit_for_manager_approval(db=db, expense_id=expense.id)
        ApprovalService.submit_for_hr_approval(db=db, expense_id=expense.id)
        
        # Run pre-screen checks and store flags
        _apply_pre_screen_checks(db, expense)
        
        # Notify manager about new expense
        await NotificationService.notify_expense_submitted(
            db=db,
            expense=expense,
            employee=current_user
        )
        
        # Refresh expense to ensure attachments are loaded
        db.refresh(expense)
        
        # Return response with attachments safely handled
        return {
            "id": expense.id,
            "category_id": expense.category_id,
            "transport_type_id": expense.transport_type_id,
            "amount": float(expense.amount),
            "expense_date": expense.expense_date,
            "description": expense.description,
            "status": expense.status.value if hasattr(expense.status, 'value') else str(expense.status),
            "created_at": expense.created_at,
            "updated_at": expense.updated_at,
            "user_id": expense.user_id,
            "first_name": expense.employee.first_name if expense.employee else None,
            "last_name": expense.employee.last_name if expense.employee else None,
            "attachments": [
                {
                    "id": att.id,
                    "file_name": att.file_name,
                    "file_type": att.file_type,
                    "file_size": att.file_size,
                    "uploaded_at": att.uploaded_at,
                    "file_path": att.file_path
                }
                for att in expense.attachments
            ] if expense.attachments else [],
            "policy_check_result": None
        }
    
    except HTTPException:
        raise
    except Exception as e:
        print(f"Exception: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Error submitting expense: {str(e)}"
        )

@router.get("/{expense_id}", response_model=ExpenseWithAttachments)
async def get_expense(
    expense_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Get expense details with attachments"""
    expense = ExpenseService.get_expense(db, expense_id)
    
    if not expense:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Expense not found"
        )
    
    # Check authorization
    if expense.user_id != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="You can only view your own expenses"
        )
    
    return expense

@router.get("/", response_model=List[dict])
async def list_expenses(
    status_filter: str = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """List expenses - for managers shows all, for employees shows their own"""
    from sqlalchemy.orm import joinedload
    
    # Managers (role_id=2) and HR (role_id=3) can see all expenses
    if current_user.role_id == 2 or current_user.role_id == 3:
        # Get all expenses for manager/HR approval with eager loading
        query = db.query(Expense).options(
            joinedload(Expense.employee),
            joinedload(Expense.attachments)
        )
        
        if status_filter:
            query = query.filter(Expense.status == status_filter)
        
        expenses = query.order_by(Expense.created_at.desc()).all()
    else:
        # Regular employees see only their own expenses
        expenses = ExpenseService.get_user_expenses(
            db=db,
            user_id=current_user.id,
            status=status_filter
        )
    
    # Format response to include employee name
    result = []
    for expense in expenses:
        expense_dict = {
            "id": expense.id,
            "category_id": expense.category_id,
            "transport_type_id": expense.transport_type_id,
            "amount": float(expense.amount),
            "expense_date": expense.expense_date,
            "description": expense.description,
            "status": expense.status,
            "created_at": expense.created_at,
            "updated_at": expense.updated_at,
            "user_id": expense.user_id,
            "first_name": expense.employee.first_name if expense.employee else "",
            "last_name": expense.employee.last_name if expense.employee else "",
            "attachments": [
                {
                    "id": att.id,
                    "filename": att.file_name,
                    "file_path": att.file_path,
                    "uploaded_at": att.uploaded_at
                }
                for att in expense.attachments
            ] if expense.attachments else []
        }
        result.append(expense_dict)
    
    return result

@router.post("/{expense_id}/extract-amount")
async def extract_receipt_amount(
    expense_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Extract amount from receipt attached to an expense.
    Useful for updating expenses that were created with placeholder amounts.
    """
    # Get expense
    expense = ExpenseService.get_expense(db, expense_id)
    if not expense:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Expense not found"
        )
    
    # Check authorization
    if expense.user_id != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="You can only extract amounts from your own expenses"
        )
    
    # Check if expense has attachments
    if not expense.attachments:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No receipts attached to this expense"
        )
    
    try:
        # Try to extract amount from first receipt
        attachment = expense.attachments[0]
        
        # Get the file path
        file_path = attachment.file_path
        file_type = attachment.file_name.split('.')[-1].lower()
        
        if not os.path.exists(file_path):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Receipt file not found: {attachment.file_name}"
            )
        
        # Extract amount
        extracted_amount, confidence, extraction_note = ImprovedReceiptExtractor.extract_amount(file_path, file_type)
        
        if not extracted_amount or extracted_amount <= 0:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Could not extract amount from receipt. {extraction_note}"
            )
        
        # Update expense with extracted amount
        old_amount = expense.amount
        expense.amount = extracted_amount
        expense.updated_at = datetime.utcnow()
        db.commit()
        db.refresh(expense)
        
        # Log the change
        AuditLogger.log(
            db=db,
            entity_type="expense",
            entity_id=expense.id,
            action="amount_extracted_from_receipt",
            performed_by=current_user.id,
            old_value={"amount": str(old_amount)},
            new_value={"amount": str(extracted_amount), "confidence": confidence}
        )
        
        return {
            "success": True,
            "expense_id": expense.id,
            "extracted_amount": float(extracted_amount),
            "confidence": confidence,
            "message": extraction_note,
            "expense": ExpenseResponse.from_orm(expense)
        }
    
    except HTTPException:
        raise
    except Exception as e:
        print(f"Error extracting amount: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error extracting amount from receipt: {str(e)}"
        )

@router.put("/{expense_id}", response_model=ExpenseDetailResponse)
async def update_expense(
    expense_id: int,
    expense_data: ExpenseUpdate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Update expense (only in SUBMITTED status)"""
    expense = ExpenseService.get_expense(db, expense_id)
    
    if not expense:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Expense not found"
        )
    
    if expense.user_id != current_user.id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="You can only update your own expenses"
        )
    
    updated_expense, error = ExpenseService.update_expense(
        db=db,
        expense_id=expense_id,
        expense_data=expense_data,
        performed_by=current_user.id
    )
    
    if error:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=error
        )
    
    return updated_expense

@router.get("/{expense_id}/validate-receipt", response_model=dict)
async def validate_expense_receipt(
    expense_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Validate the genuineness of receipts for an expense.
    Returns detailed validation results including risk factors and recommendations.
    """
    # Get expense with attachments
    expense = db.query(Expense).filter(Expense.id == expense_id).first()
    if not expense:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Expense not found"
        )
    
    # Check authorization
    if expense.user_id != current_user.id and current_user.role not in ['admin', 'manager']:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not authorized to validate this expense"
        )
    
    # Get attachments
    attachments = db.query(ExpenseAttachment).filter(
        ExpenseAttachment.expense_id == expense_id
    ).all()
    
    if not attachments:
        return {
            "message": "No receipts found for validation",
            "validation_results": None
        }
    
    validation_service = ReceiptValidationService()
    all_validation_results = []
    
    for attachment in attachments:
        try:
            # Get file path
            file_path = attachment.file_path
            
            # Extract text based on file type
            file_extension = attachment.filename.split('.')[-1].lower()
            full_text = ""
            
            if file_extension == 'pdf':
                try:
                    import pdfplumber
                    with pdfplumber.open(file_path) as pdf:
                        for page in pdf.pages:
                            full_text += (page.extract_text() or "") + "\n"
                except Exception as e:
                    full_text = f"Error extracting PDF text: {str(e)}"
            
            elif file_extension in ['docx', 'doc']:
                try:
                    from docx import Document
                    doc = Document(file_path)
                    for paragraph in doc.paragraphs:
                        full_text += paragraph.text + " "
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                full_text += cell.text + " "
                except Exception as e:
                    full_text = f"Error extracting Word text: {str(e)}"
            
            # Validate receipt
            validation_results = validation_service.validate_receipt(
                file_path, full_text, expense.amount
            )
            
            # Add attachment info
            validation_results['attachment_id'] = attachment.id
            validation_results['filename'] = attachment.filename
            validation_results['file_type'] = file_extension
            
            all_validation_results.append(validation_results)
            
        except Exception as e:
            all_validation_results.append({
                "attachment_id": attachment.id,
                "filename": attachment.filename,
                "error": f"Validation failed: {str(e)}",
                "is_genuine": False,
                "confidence_score": 0,
                "risk_level": "Error"
            })
    
    # Overall assessment
    if all_validation_results:
        avg_confidence = sum(r.get('confidence_score', 0) for r in all_validation_results) / len(all_validation_results)
        overall_genuine = all(r.get('is_genuine', False) for r in all_validation_results)
        
        # Determine overall risk level
        if avg_confidence >= 80:
            overall_risk = "Low"
        elif avg_confidence >= 60:
            overall_risk = "Medium"
        else:
            overall_risk = "High"
    else:
        avg_confidence = 0
        overall_genuine = False
        overall_risk = "Unknown"
    
    return {
        "expense_id": expense_id,
        "expense_amount": expense.amount,
        "overall_assessment": {
            "is_genuine": overall_genuine,
            "average_confidence": avg_confidence,
            "risk_level": overall_risk,
            "total_receipts_validated": len(all_validation_results)
        },
        "individual_validations": all_validation_results,
        "recommendations": [
            "Review receipts with low confidence scores manually",
            "Request additional documentation for high-risk expenses",
            "Verify vendor details for large amounts"
        ]
    }

@router.get("/receipts/{attachment_id}")
async def get_receipt(
    attachment_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Get receipt file for viewing
    
    Returns the receipt file for the given attachment ID.
    Only the expense owner, managers, and finance can view.
    """
    try:
        # Get the attachment
        attachment = db.query(ExpenseAttachment).filter(
            ExpenseAttachment.id == attachment_id
        ).first()
        
        if not attachment:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Receipt not found"
            )
        
        # Get the expense to check authorization
        expense = db.query(Expense).filter(Expense.id == attachment.expense_id).first()
        if not expense:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Expense not found"
            )
        
        # Check if user is authorized to view
        # Owner, managers (role 2), and finance (role 3) can view
        is_owner = expense.user_id == current_user.id
        is_manager_or_finance = current_user.role_id in [2, 3]
        
        if not (is_owner or is_manager_or_finance):
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Not authorized to view this receipt"
            )
        
        # Return receipt file info and URL
        return {
            "attachment_id": attachment.id,
            "filename": attachment.file_name,
            "file_path": f"/api/expenses/file/{attachment.file_path}",
            "uploaded_at": attachment.uploaded_at
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error retrieving receipt: {str(e)}"
        )

@router.get("/file/{file_path:path}")
async def get_file(
    file_path: str,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Serve the actual receipt file
    
    Returns the file content for download/viewing.
    """
    try:
        from fastapi.responses import FileResponse
        import mimetypes
        
        # Construct full file path
        full_path = os.path.join("/app/bills", file_path)
        
        # Security check: ensure the path is within the bills directory
        if not os.path.abspath(full_path).startswith("/app/bills"):
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Invalid file path"
            )
        
        if not os.path.exists(full_path):
            logger.error(f"File not found: {full_path}")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"File not found: {file_path}"
            )
        
        # Determine media type
        media_type, _ = mimetypes.guess_type(full_path)
        if not media_type:
            media_type = "application/octet-stream"
        
        logger.info(f"Serving file: {full_path} with media type: {media_type}")
        
        return FileResponse(full_path, media_type=media_type)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error serving file: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error serving file: {str(e)}"
        )


def _apply_pre_screen_checks(db: Session, expense: Expense):
    """
    Compute and store rule-based pre-screen flags and recommendation.
    This does not call Llama; it only uses deterministic rules.
    """
    flags = []
    score = 100.0

    # Date flag (already enforced on submit, but store for visibility)
    today = datetime.utcnow().date()
    if expense.expense_date:
        days_old = (today - expense.expense_date).days
        if days_old > 31:
            flags.append("DATE_TOO_OLD")
            score -= 30.0
        elif days_old < 0:
            flags.append("DATE_IN_FUTURE")
            score -= 30.0

    # Policy violation flag
    if expense.policy_check_result:
        try:
            policy_data = isinstance(expense.policy_check_result, dict) and expense.policy_check_result or {}
            violations = policy_data.get("violations")
            if isinstance(violations, list) and violations:
                flags.append("POLICY_VIOLATION")
                score -= 20.0
        except Exception:
            pass

    # Receipt quality flag
    if not expense.extracted_text or len(expense.extracted_text.strip()) < 10:
        flags.append("NO_TEXT_EXTRACTED")
        score -= 25.0

    # Amount sanity
    try:
        amt = float(expense.amount)
        if amt <= 0:
            flags.append("AMOUNT_MISSING_OR_ZERO")
            score -= 20.0
    except Exception:
        flags.append("AMOUNT_MISSING_OR_ZERO")
        score -= 20.0

    score = max(0.0, min(100.0, score))

    # Derive recommendation
    if score >= 80:
        recommendation = "✅ SAFE TO APPROVE"
    elif score >= 60:
        recommendation = "⚠️ NEEDS REVIEW"
    else:
        recommendation = "❌ RECOMMEND REJECTION"

    # Store on expense
    expense.validation_score = score
    expense.risk_factors = {"flags": flags}
    expense.ai_analysis = {
        "source": "pre_screen",
        "recommendation": recommendation,
        "flags": flags,
        "score": score
    }
    db.commit()
