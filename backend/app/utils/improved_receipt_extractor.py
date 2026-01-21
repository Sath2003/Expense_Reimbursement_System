import re
import os
from typing import Optional, Tuple, List
from pathlib import Path
from decimal import Decimal

try:
    import pdfplumber
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from PIL import Image
    IMAGE_SUPPORT = True
except ImportError:
    IMAGE_SUPPORT = False

try:
    import pytesseract
    import os
    import platform
    PYTESSERACT_SUPPORT = True
    
    # Auto-configure Tesseract path based on OS
    if platform.system() == 'Windows':
        # Try common Windows installation paths
        tesseract_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        ]
        for path in tesseract_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.pytesseract_cmd = path
                break
    # For Linux/Mac, it should be in PATH automatically
except ImportError:
    PYTESSERACT_SUPPORT = False
except Exception as e:
    PYTESSERACT_SUPPORT = False

try:
    from openpyxl import load_workbook
    EXCEL_SUPPORT = True
except ImportError:
    EXCEL_SUPPORT = False

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    if platform.system() == 'Windows':
        # Try common Windows installation paths
        tesseract_paths = [
            r'C:\Program Files\Tesseract-OCR\tesseract.exe',
            r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        ]
        for path in tesseract_paths:
            if os.path.exists(path):
                pytesseract.pytesseract.pytesseract_cmd = path
                break
    # For Linux/Mac, it should be in PATH automatically
except ImportError:
    PYTESSERACT_SUPPORT = False
except Exception as e:
    PYTESSERACT_SUPPORT = False
    print(f"Warning: Could not configure pytesseract: {e}")


class ImprovedReceiptExtractor:
    """
    Enhanced utility to extract amount information from receipt files.
    Supports PDF and image files with multiple fallback strategies.
    """
    
    # Comprehensive patterns for Indian currency amounts
    AMOUNT_PATTERNS = [
        # TIER 1: EXPLICIT RUPEE SYMBOL - HIGHEST PRIORITY (must have ₹ or Rs)
        (r'₹\s*([\d,]+\.\d{1,2})', 'rupee_symbol_decimal'),  # ₹156.00 / ₹156.0
        (r'₹\s*([1-9]\d{0,6}(?:\.\d{1,2})?)', 'rupee_symbol'),  # ₹156
        (r'Rs\.?\s*([\d,]+\.\d{1,2})', 'rs_abbrev_decimal'),  # Rs156.00 / Rs156.0
        (r'Rs\.?\s*([1-9]\d{0,6}(?:\.\d{1,2})?)', 'rs_abbrev'),  # Rs156
        (r'INR\s*([\d,]+\.\d{1,2})', 'inr_code_decimal'),  # INR 156.00 / INR 156.0
        (r'INR\s*([1-9]\d{0,6}(?:\.\d{1,2})?)', 'inr_code'),  # INR 156
        
        # TIER 2: TOTAL/AMOUNT KEYWORDS WITH RUPEE (must have keyword + currency indicator)
        (r'(?:total\s+(?:amount|paid|fare|price)|amount\s+paid|final\s+amount|grand\s+total|net\s+amount|bill\s+amount|sub\s+total)[\s:]*₹\s*([\d,]+\.?\d*)', 'total_rupee_final'),
        (r'(?:total|amount|price|cost|due|payable|net|fare|bill)[\s:]*₹\s*([\d,]+\.?\d*)', 'total_rupee'),

        # TIER 2B: TOTAL KEYWORDS WITHOUT CURRENCY (common in tables like "Total 1349")
        (r'(?:grand\s+total|net\s+total|total)\s*[:\-]?\s*([1-9]\d{0,6}(?:\.\d{1,2})?)', 'total_no_currency'),
        
        # TIER 3: TEXT KEYWORDS + CURRENCY WORDS (very explicit)
        (r'(?:total|amount|price|cost|due|payable|fare|paid|bill|sum|charge)\s+(?:is|:|=)?\s*(?:Rs|rupees?|INR)\s*[\.]?\s*([\d,]+\.\d{1,2})', 'text_total'),
        (r'rupees?\s+([\d,]+\.\d{1,2})', 'rupees_text'),
        
        # TIER 4: AMOUNTS STRICTLY WITH CURRENCY SYMBOL NEARBY
        (r'([\d,]+\.\d{1,2})\s*(?:Rs|INR|₹)', 'table_amount'),
        
        # NOTE: All patterns without explicit currency context removed
        # This prevents UPI IDs and other numbers from being extracted
    ]
    
    @staticmethod
    def _normalize_amount(amount_str: str) -> Optional[float]:
        """
        Convert extracted amount string to float.
        Handles various formats: 1,00,000.50 | 100,000.50 | 1000.50
        Filters out year-like patterns (19XX, 20XX) and UPI IDs/phone numbers.
        """
        if not amount_str:
            return None
        
        amount_str = amount_str.strip()
        
        # Remove common unwanted characters but keep decimal points and commas
        amount_str = re.sub(r'[^\d.,]', '', amount_str)
        
        # Handle empty string after cleanup
        if not amount_str:
            return None
        
        # CRITICAL: Filter out UPI IDs and phone numbers
        # - UPI IDs: 10-11 digit numbers without decimal points
        # - Phone numbers: 10 digits (Indian mobile)
        # - Partial UPI IDs: 6-9 digit substrings from UPI IDs without decimal
        clean_digits = amount_str.replace(',', '')
        
        # Reject if it looks like an undecorated number without decimal (likely UPI/phone)
        if '.' not in amount_str:
            # If more than 5 digits and no decimal point, very likely UPI/phone/ID number
            if len(clean_digits) > 5:
                return None
        
        # If decimal exists, accept 1-2 decimals (OCR often misses a digit)
        if '.' in amount_str:
            parts = amount_str.split('.')
            if len(parts) != 2:
                return None
            decimals = parts[1]
            if len(decimals) == 1:
                amount_str = f"{parts[0]}.{decimals}0"
            elif len(decimals) > 2:
                amount_str = f"{parts[0]}.{decimals[:2]}"
        
        # Handle Indian numbering (1,00,000) vs Western (100,000)
        if ',' in amount_str:
            # Split by comma to analyze
            parts = amount_str.split(',')
            
            # If there's a decimal point, handle it specially
            if '.' in parts[-1]:
                last_part = parts[-1]
                if len(last_part.split('.')[0]) == 2:  # 2 digits before decimal = Indian format
                    amount_str = ''.join(parts)
                else:  # Western format
                    amount_str = amount_str.replace(',', '')
            else:
                # No decimal - check pattern
                if len(parts) == 3 and len(parts[-1]) == 3:  # 1,23,456 format (Indian)
                    amount_str = ''.join(parts)
                elif len(parts) == 2 and len(parts[-1]) == 3:  # 1,234 format (Western)
                    amount_str = amount_str.replace(',', '')
                elif len(parts[-1]) == 2:  # Last part has 2 digits = likely Indian format
                    amount_str = ''.join(parts)
                else:
                    amount_str = amount_str.replace(',', '')
        
        try:
            value = float(amount_str)
            # Enhanced validation: amount should be reasonable
            if 0.01 <= value <= 999999.99:
                # Less restrictive year filtering - only filter obvious years
                if not (1900 <= value <= 2099 and len(str(int(value))) == 4 and '.' not in amount_str):
                    return value
        except ValueError:
            pass
        
        return None
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> Tuple[Optional[float], str, str]:
        """
        Extract amount from PDF file using pdfplumber.
        Returns: (amount, confidence, status_message)
        """
        if not PDF_SUPPORT:
            return None, "none", "PDF library not installed"
        
        if not os.path.exists(file_path):
            return None, "none", f"File not found: {file_path}"
        
        try:
            amounts_with_patterns: List[Tuple[float, str]] = []
            
            with pdfplumber.open(file_path) as pdf:
                full_text = ""
                page_count = 0
                table_totals: List[float] = []
                
                for page in pdf.pages:
                    page_count += 1
                    # Extract text with better formatting
                    page_text = page.extract_text() or ""
                    full_text += page_text + "\n"
                    
                    # Also try to extract tables separately for better amount detection
                    try:
                        tables = page.extract_tables()
                        for table in tables:
                            for row in table:
                                row_text = " ".join([c for c in row if isinstance(c, str)])
                                if row_text and re.search(r'\b(total|grand\s+total|net\s+total)\b', row_text, re.IGNORECASE):
                                    for cell in row:
                                        if cell and isinstance(cell, str):
                                            match = re.search(r'([1-9]\d{0,6}(?:\.\d{2})?)', cell.replace(',', ''))
                                            if match:
                                                try:
                                                    value = float(match.group(1))
                                                    if value >= 10:
                                                        table_totals.append(value)
                                                except ValueError:
                                                    pass
                                for cell in row:
                                    if cell and isinstance(cell, str):
                                        full_text += cell + " "
                    except:
                        pass

                if table_totals:
                    final_amount = max(table_totals)
                    return final_amount, "high", f"PDF table total detected: ₹{final_amount:.2f}"
                
                # If no text extracted, try OCR as fallback
                if not full_text.strip() and PYTESSERACT_SUPPORT:
                    try:
                        from PIL import Image
                        import io
                        
                        # Convert PDF page to image for OCR
                        for i, page in enumerate(pdf.pages[:3]):  # Try first 3 pages
                            try:
                                # Get page as image
                                img = page.to_image(resolution=150)
                                if img:
                                    # OCR the image
                                    import pytesseract
                                    ocr_text = pytesseract.image_to_string(img)
                                    if ocr_text and len(ocr_text.strip()) > 10:
                                        full_text += f"\n[OCR_PAGE_{i+1}] {ocr_text}"
                            except:
                                continue
                    except Exception as ocr_error:
                        print(f"[DEBUG] OCR fallback failed: {ocr_error}")
            
            # Clean up the extracted text
            full_text = re.sub(r'\s+', ' ', full_text)  # Normalize whitespace
            full_text = full_text.strip()
            
            # Debug: Print extracted text and metadata
            print(f"[DEBUG] PDF processed: {page_count} pages")
            print(f"[DEBUG] PDF extracted text: {full_text[:800]}...")
            print(f"[DEBUG] Text length: {len(full_text)} characters")
            
            # Search using all patterns
            for pattern, pattern_name in ImprovedReceiptExtractor.AMOUNT_PATTERNS:
                matches = re.finditer(pattern, full_text, re.IGNORECASE)
                for match in matches:
                    amount_str = match.group(1)
                    print(f"[DEBUG] Found match: pattern={pattern_name}, amount_str={amount_str}")
                    amount = ImprovedReceiptExtractor._normalize_amount(amount_str)
                    if amount and amount >= 10:  # Filter out very small numbers
                        # Additional filter: reject year-like amounts (1900-2099) from fallback pattern
                        # This filters out dates that match the fallback pattern
                        if not (1900 <= amount <= 2099 and pattern_name == 'fallback_number'):
                            print(f"[DEBUG] Valid amount: {amount} from pattern {pattern_name}")
                            amounts_with_patterns.append((amount, pattern_name))
            
            if amounts_with_patterns:
                # Priority 1: Most specific patterns
                final_patterns = [a for a, p in amounts_with_patterns if p == 'total_rupee_final']
                high_priority = [a for a, p in amounts_with_patterns if p in ['total_rupee', 'text_total', 'generic_total']]
                currency_priority = [a for a, p in amounts_with_patterns if p in ['rupee_symbol', 'rs_abbrev', 'inr_code', 'rupees_text', 'currency_context', 'table_amount']]
                decimal_priority = [a for a, p in amounts_with_patterns if p in ['decimal_amount', 'large_number', 'amt_pattern', 'pay_pattern']]
                
                if final_patterns:
                    # Use the most specific pattern
                    final_amount = max(final_patterns)
                    confidence = "high"
                elif high_priority:
                    # Among high priority amounts, pick the largest (final total)
                    final_amount = max(high_priority)
                    confidence = "high"
                elif currency_priority:
                    # Use amounts with clear currency indicators
                    final_amount = max(currency_priority)
                    confidence = "medium"
                elif decimal_priority:
                    # Use amounts with decimal places (likely real amounts)
                    final_amount = max(decimal_priority)
                    confidence = "low"
                else:
                    # Fall back to largest amount found
                    final_amount = max(amounts_with_patterns, key=lambda x: x[0])[0]
                    confidence = "very_low"
                
                return final_amount, confidence, f"PDF extraction successful: ₹{final_amount:.2f}"
            
            return None, "low", "No amount pattern found in PDF"
        
        except Exception as e:
            return None, "none", f"PDF extraction error: {str(e)}"
    
    @staticmethod
    def extract_from_image(file_path: str) -> Tuple[Optional[float], str, str]:
        """
        Extract amount from image using OCR (if available).
        Without pytesseract, image extraction is not supported.
        
        Returns: (amount, confidence, status_message)
        """
        if not IMAGE_SUPPORT:
            return None, "none", "Image library not installed"
        
        if not os.path.exists(file_path):
            return None, "none", f"File not found: {file_path}"
        
        if not PYTESSERACT_SUPPORT:
            return None, "none", "OCR not available - pytesseract not installed. Install with: pip install pytesseract tesseract-ocr"
        
        try:
            from PIL import Image, ImageEnhance
            
            amounts_with_patterns: List[Tuple[float, str]] = []
            
            # Open and preprocess image
            image = Image.open(file_path)
            
            try:
                # First attempt: direct OCR
                text = pytesseract.image_to_string(image)
                
                # Second attempt: enhanced image if first didn't work
                if not text or len(text.strip()) < 10:
                    # Enhance contrast and convert to grayscale
                    enhancer = ImageEnhance.Contrast(image.convert('L'))
                    enhanced = enhancer.enhance(2)
                    text = pytesseract.image_to_string(enhanced)

                # Normalize OCR text (common OCR misreads on receipts)
                text = text.replace("\u00a0", " ")
                text = re.sub(r"\s+", " ", text)
                # Some receipts OCR ₹ as 'n' or similar. Map a standalone 'n' before digits to ₹.
                text = re.sub(r"(?<!\w)n\s*(?=\d)", "₹", text, flags=re.IGNORECASE)
                
                # Search patterns in OCR text
                for pattern, pattern_name in ImprovedReceiptExtractor.AMOUNT_PATTERNS:
                    matches = re.finditer(pattern, text, re.IGNORECASE)
                    for match in matches:
                        amount_str = match.group(1)
                        amount = ImprovedReceiptExtractor._normalize_amount(amount_str)
                        if amount and amount >= 10:
                            amounts_with_patterns.append((amount, pattern_name))
            
            except Exception as ocr_err:
                return None, "low", f"OCR processing failed: {str(ocr_err)}. Make sure Tesseract is installed."
            
            if amounts_with_patterns:
                # Priority 1: Most specific patterns
                final_patterns = [a for a, p in amounts_with_patterns if p == 'total_rupee_final']
                high_priority = [a for a, p in amounts_with_patterns if p in ['total_rupee', 'text_total', 'generic_total']]
                currency_priority = [a for a, p in amounts_with_patterns if p in ['rupee_symbol', 'rs_abbrev', 'inr_code', 'rupees_text', 'currency_context', 'table_amount']]
                decimal_priority = [a for a, p in amounts_with_patterns if p in ['decimal_amount', 'large_number', 'amt_pattern', 'pay_pattern']]
                
                if final_patterns:
                    final_amount = max(final_patterns)
                    confidence = "high"
                elif high_priority:
                    final_amount = max(high_priority)
                    confidence = "high"
                elif currency_priority:
                    final_amount = max(currency_priority)
                    confidence = "medium"
                elif decimal_priority:
                    final_amount = max(decimal_priority)
                    confidence = "low"
                else:
                    final_amount = max(amounts_with_patterns, key=lambda x: x[0])[0]
                    confidence = "very_low"
                
                return final_amount, confidence, f"Image OCR successful: ₹{final_amount:.2f}"
            
            return None, "low", "No amount pattern found in image OCR"
        
        except Exception as e:
            return None, "none", f"Image extraction error: {str(e)}"
    
    @staticmethod
    def extract_from_excel(file_path: str) -> Tuple[Optional[float], str, str]:
        """
        Extract amount from Excel file (.xlsx, .xls) using openpyxl.
        
        Returns: (amount, confidence, status_message)
        """
        if not EXCEL_SUPPORT:
            return None, "none", "Excel library not installed (openpyxl required)"
        
        if not os.path.exists(file_path):
            return None, "none", f"File not found: {file_path}"
        
        try:
            amounts_with_patterns: List[Tuple[float, str]] = []
            
            workbook = load_workbook(file_path, data_only=True)
            full_text = ""
            
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]
                for row in worksheet.iter_rows(values_only=True):
                    for value in row:
                        if value is None:
                            continue
                        full_text += str(value) + " "
            
            full_text = re.sub(r'\s+', ' ', full_text).strip()
            
            for pattern, pattern_name in ImprovedReceiptExtractor.AMOUNT_PATTERNS:
                matches = re.finditer(pattern, full_text, re.IGNORECASE)
                for match in matches:
                    amount_str = match.group(1)
                    amount = ImprovedReceiptExtractor._normalize_amount(amount_str)
                    if amount and amount >= 10:
                        amounts_with_patterns.append((amount, pattern_name))
            
            if amounts_with_patterns:
                final_patterns = [a for a, p in amounts_with_patterns if p == 'total_rupee_final']
                high_priority = [a for a, p in amounts_with_patterns if p in ['total_rupee', 'text_total', 'generic_total']]
                currency_priority = [a for a, p in amounts_with_patterns if p in ['rupee_symbol', 'rs_abbrev', 'inr_code', 'rupees_text', 'currency_context', 'table_amount']]
                decimal_priority = [a for a, p in amounts_with_patterns if p in ['decimal_amount', 'large_number', 'amt_pattern', 'pay_pattern']]
                
                if final_patterns:
                    final_amount = max(final_patterns)
                    confidence = "high"
                elif high_priority:
                    final_amount = max(high_priority)
                    confidence = "high"
                elif currency_priority:
                    final_amount = max(currency_priority)
                    confidence = "medium"
                elif decimal_priority:
                    final_amount = max(decimal_priority)
                    confidence = "low"
                else:
                    final_amount = max(amounts_with_patterns, key=lambda x: x[0])[0]
                    confidence = "very_low"
                
                return final_amount, confidence, f"Excel extraction successful: ₹{final_amount:.2f}"
            
            return None, "low", "No amount pattern found in Excel"
        
        except Exception as e:
            return None, "none", f"Excel extraction error: {str(e)}"
    
    @staticmethod
    def extract_from_word(file_path: str) -> Tuple[Optional[float], str, str]:
        """
        Extract amount from Word document (.docx) using python-docx.
        
        Returns: (amount, confidence, status_message)
        """
        if not DOCX_SUPPORT:
            return None, "none", "Word document library not installed (python-docx required)"
        
        if not os.path.exists(file_path):
            return None, "none", f"File not found: {file_path}"
        
        try:
            amounts_with_patterns: List[Tuple[float, str]] = []
            
            doc = Document(file_path)
            full_text = ""
            
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + " "
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += (cell.text or "") + " "
            
            full_text = re.sub(r'\s+', ' ', full_text).strip()
            
            for pattern, pattern_name in ImprovedReceiptExtractor.AMOUNT_PATTERNS:
                matches = re.finditer(pattern, full_text, re.IGNORECASE)
                for match in matches:
                    amount_str = match.group(1)
                    amount = ImprovedReceiptExtractor._normalize_amount(amount_str)
                    if amount and amount >= 10:
                        amounts_with_patterns.append((amount, pattern_name))
            
            if amounts_with_patterns:
                final_patterns = [a for a, p in amounts_with_patterns if p == 'total_rupee_final']
                high_priority = [a for a, p in amounts_with_patterns if p in ['total_rupee', 'text_total', 'generic_total']]
                currency_priority = [a for a, p in amounts_with_patterns if p in ['rupee_symbol', 'rs_abbrev', 'inr_code', 'rupees_text', 'currency_context', 'table_amount']]
                decimal_priority = [a for a, p in amounts_with_patterns if p in ['decimal_amount', 'large_number', 'amt_pattern', 'pay_pattern']]
                
                if final_patterns:
                    final_amount = max(final_patterns)
                    confidence = "high"
                elif high_priority:
                    final_amount = max(high_priority)
                    confidence = "high"
                elif currency_priority:
                    final_amount = max(currency_priority)
                    confidence = "medium"
                elif decimal_priority:
                    final_amount = max(decimal_priority)
                    confidence = "low"
                else:
                    final_amount = max(amounts_with_patterns, key=lambda x: x[0])[0]
                    confidence = "very_low"
                
                return final_amount, confidence, f"Word extraction successful: ₹{final_amount:.2f}"
            
            return None, "low", "No amount pattern found in Word document"
        
        except Exception as e:
            return None, "none", f"Word extraction error: {str(e)}"
    
    @staticmethod
    def extract_amount(file_path: str, file_type: str) -> Tuple[Optional[float], str, str]:
        """
        Extract amount from receipt file (PDF, Excel, Word, or image).
        
        Supported formats:
        - PDF: .pdf
        - Excel: .xlsx, .xls
        - Word: .docx, .doc
        - Images: .jpg, .jpeg, .png, .gif, .bmp, .tiff, .webp
        
        Args:
            file_path: Full path to the receipt file
            file_type: File extension (pdf, xlsx, docx, jpg, png, etc.)
        
        Returns:
            (extracted_amount, confidence_level, status_message)
            confidence_level: 'high', 'medium', 'low', 'none'
        """
        file_type = file_type.lower()
        
        if file_type == 'pdf':
            return ImprovedReceiptExtractor.extract_from_pdf(file_path)
        elif file_type in ['xlsx', 'xls']:
            return ImprovedReceiptExtractor.extract_from_excel(file_path)
        elif file_type in ['docx', 'doc']:
            return ImprovedReceiptExtractor.extract_from_word(file_path)
        elif file_type in ['jpg', 'jpeg', 'png', 'gif', 'bmp', 'tiff', 'webp']:
            return ImprovedReceiptExtractor.extract_from_image(file_path)
        else:
            return None, "none", f"Unsupported file type: {file_type}. Supported: PDF, Excel (xlsx/xls), Word (docx/doc), Images (jpg/png/gif/bmp/tiff/webp)"
    
    @staticmethod
    def check_capabilities() -> dict:
        """
        Check what extraction capabilities are available.
        
        Returns:
            dict with status of each method
        """
        return {
            "pdf_support": PDF_SUPPORT,
            "image_support": IMAGE_SUPPORT,
            "ocr_support": PYTESSERACT_SUPPORT,
            "excel_support": EXCEL_SUPPORT,
            "word_support": DOCX_SUPPORT,
            "pdf_library": "pdfplumber" if PDF_SUPPORT else "not installed",
            "image_library": "PIL/Pillow" if IMAGE_SUPPORT else "not installed",
            "ocr_library": "pytesseract" if PYTESSERACT_SUPPORT else "not installed",
            "excel_library": "openpyxl" if EXCEL_SUPPORT else "not installed",
            "word_library": "python-docx" if DOCX_SUPPORT else "not installed",
        }
